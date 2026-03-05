"""Native DOCX parser using python-docx."""
from __future__ import annotations

import io
from dataclasses import dataclass

from docx import Document
from app.core.logging_config import get_logger

logger = get_logger(__name__)


@dataclass
class ParseResult:
    text: str
    confidence: float
    method: str = "primary"
    error: str | None = None


def parse_docx(file_bytes: bytes, file_name: str = "") -> ParseResult:
    """
    Extract text from a DOCX file using python-docx.
    Reads paragraphs and table cells for comprehensive text extraction.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in parts:
                        parts.append(text)

        full_text = "\n".join(parts)
        confidence = 1.0 if parts else 0.0

        logger.info(
            "docx_parsed",
            file=file_name,
            paragraphs=len(doc.paragraphs),
            chars=len(full_text),
        )
        return ParseResult(text=full_text, confidence=confidence, method="primary")

    except Exception as e:
        logger.error("docx_parse_failed", file=file_name, error=str(e))
        return ParseResult(text="", confidence=0.0, method="primary", error=str(e))
